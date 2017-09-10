import json
import core.common.utils as utils
import os
from docx import Document
from datetime import datetime
import locale

FILES_TO_SAVE = "c:\\AssessoriaManager\\files"

class DocumentBo():
    def __init__(self):
        locale.setlocale(locale.LC_ALL, 'ptb_bra')

    def get_document(self, file, cliente):
        filename = "contrato.docx" if file == "contrato" else "richiedente.docx"
        documents_directory = os.path.join(os.getcwd(), "documents")
        cliente_directory = utils.create_directory(FILES_TO_SAVE, cliente["cognome"] + " " + cliente["nome"])
        document_to_edit = os.path.join(documents_directory, filename)
        filename_to_download = self.edit_document(document_to_edit, cliente_directory, cliente, file)
        return os.path.join(cliente_directory, filename_to_download)

    def edit_document(self, document_to_edit, cliente_directory, cliente, file):
        if file == "contrato":
            return self.edit_contrato(document_to_edit, cliente_directory, cliente)
        else:
            return self.edit_richiedente(document_to_edit, cliente_directory, cliente)

    def edit_contrato(self, document_to_edit, cliente_directory, cliente):
        document = Document(document_to_edit)
        new_name = "contrato_" + cliente["cognome"] + "_" + cliente["nome"] + ".docx"
        for paragraph in document.paragraphs:
            inline = paragraph.runs
            for i in range(len(inline)):
                if '{{cliente[“residencia”]}}' in inline[i].text:
                    text = inline[i].text.replace('{{cliente[“residencia”]}}', cliente["residencia"])
                    inline[i].text = text
                if '{{cliente[“nome”] cliente[“cognome”]}}' in inline[i].text:
                    text = inline[i].text.replace('{{cliente[“nome”] cliente[“cognome”]}}', cliente["nome"] + " " + cliente["cognome"])
                    inline[i].text = text
                if '{{cliente[“estado_civil”]}}' in inline[i].text:
                    text = inline[i].text.replace('{{cliente[“estado_civil”]}}', cliente["estado_civil"])
                    inline[i].text = text
                if '{{data}}' in inline[i].text:
                    text = inline[i].text.replace('{{data}}', "{0:%d} de {0:%B} de {0:%Y}".format(datetime.now().date()))
                    inline[i].text = text
                if '{{cliente[“passaporte”]}}' in inline[i].text:
                    text = inline[i].text.replace('{{cliente[“passaporte”]}}', cliente["passaporte"])
                    inline[i].text = text
        document.save(os.path.join(cliente_directory, new_name))
        return new_name

    def edit_richiedente(self, document_to_edit, cliente_directory, cliente):
        document = Document(document_to_edit)
        new_name = "richiedente_" + cliente["cognome"] + "_" + cliente["nome"] + ".docx"
        for paragraph in document.paragraphs:
            inline = paragraph.runs
            for i in range(len(inline)):
                if '{{cliente[“cognome”]}}' in inline[i].text:
                    text = inline[i].text.replace('{{cliente[“cognome”]}}', cliente["cognome"])
                    inline[i].text = text
                if '{{cliente[“nome”]}}' in inline[i].text:
                    text = inline[i].text.replace('{{cliente[“nome”]}}', cliente["nome"])
                    inline[i].text = text
                if '{{cliente[“estado_civil”]}}' in inline[i].text:
                    text = inline[i].text.replace('{{cliente[“estado_civil”]}}', cliente["estado_civil"])
                    inline[i].text = text
                if '{{cliente[“email”]}}' in inline[i].text:
                    text = inline[i].text.replace('{{cliente[“email”]}}', cliente["email"])
                    inline[i].text = text
                if '{{cliente[“residencia”]}}' in inline[i].text:
                    text = inline[i].text.replace('{{cliente[“residencia”]}}', cliente["residencia"])
                    inline[i].text = text
                if '{{cliente[“nato_a”]}}' in inline[i].text:
                    text = inline[i].text.replace('{{cliente[“nato_a”]}}', cliente["nato_a"])
                    inline[i].text = text
                if '{{cliente[“data_nascimento”]}}' in inline[i].text:
                    text = inline[i].text.replace('{{cliente[“data_nascimento”]}}', utils.get_data(cliente["data_nascimento"]))
                    inline[i].text = text
                if '{{cliente[“formacao”]}}' in inline[i].text:
                    text = inline[i].text.replace('{{cliente[“formacao”]}}', get_formacao(cliente["formacao"]))
                    inline[i].text = text
                if '{{cliente[“profissao”]}}' in inline[i].text:
                    text = inline[i].text.replace('{{cliente[“profissao”]}}', cliente["profissao"])
                    inline[i].text = text
                if '{{cliente[“cognome_pai”]}}' in inline[i].text:
                    text = inline[i].text.replace('{{cliente[“cognome_pai”]}}', cliente["cognome_pai"])
                    inline[i].text = text
                if '{{cliente[“nome_pai”]}}' in inline[i].text:
                    text = inline[i].text.replace('{{cliente[“nome_pai”]}}', cliente["nome_pai"])
                    inline[i].text = text
                if '{{cliente[“nato_a_pai”]}}' in inline[i].text:
                    text = inline[i].text.replace('{{cliente[“nato_a_pai”]}}', cliente["nato_a_pai"])
                    inline[i].text = text
                if '{{cliente[“data_nascimento_pai”]}}' in inline[i].text:
                    text = inline[i].text.replace('{{cliente[“data_nascimento_pai”]}}', utils.get_data(cliente["data_nascimento_pai"]))
                    inline[i].text = text
                if '{{cliente[“cognome_mae”]}}' in inline[i].text:
                    text = inline[i].text.replace('{{cliente[“cognome_mae”]}}', cliente["cognome_mae"])
                    inline[i].text = text
                if '{{cliente[“nome_mae”]}}' in inline[i].text:
                    text = inline[i].text.replace('{{cliente[“nome_mae”]}}', cliente["nome_mae"])
                    inline[i].text = text
                if '{{cliente[“nato_a_mae”]}}' in inline[i].text:
                    text = inline[i].text.replace('{{cliente[“nato_a_mae”]}}', cliente["nato_a_mae"])
                    inline[i].text = text
                if '{{cliente[“data_nascimento_mae”]}}' in inline[i].text:
                    text = inline[i].text.replace('{{cliente[“data_nascimento_mae”]}}', utils.get_data(cliente["data_nascimento_mae"]))
                    inline[i].text = text
        document.save(os.path.join(cliente_directory, new_name))
        return new_name

    def save_document(self, file, entity):
        try:
            directory = utils.create_directory(FILES_TO_SAVE, entity["cognome"] + " " + entity["nome"])
            file.save(os.path.join(directory, file.filename))
            return True
        except:
            return False

def get_formacao(str_formacao):
    if str_formacao == "superiorc":
        return "Superior Completo"
    elif str_formacao == "superiori":
        return "Superior Incompleto"
    elif str_formacao == "medioc":
        return "Ensino Médio Completo"
    elif str_formacao == "medioi":
        return "Ensino Médio Incompleto"
    elif str_formacao == "fundamentalc":
        return "Fundamental Completo"
    elif str_formacao == "fundamentali":
        return "Fundamental Incompleto"
    else:
        return ""